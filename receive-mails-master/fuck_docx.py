# coding:utf-8
import docx
from docx.document import Document as dc
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor  # 设置字体颜色
from docx.oxml.ns import qn  # 设置中文字体
import pandas as pd

FILE_PATH = r"/Users/lowenve/Downloads/QQ_FileRecv/20200428.docx"

obj = docx.Document(FILE_PATH)


def iter_block_items(parent):
    # print('utils.py ----> iter_block_items:', 2)
    if isinstance(parent, dc):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("[TypeError] Document in insuitable type.")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table2list(table):
    data = []
    for i, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        data.append(row_data)
    return data


# 替换的段落关键字
word = '企业'
# 替换的表格关键字
table_text = '表格关键字'


def set_run(run, font_size, bold, color, name):
    '''
    设置run对象
    :param run:
    :param font_size: 字体大小
    :param bold: 是否加粗
    :param color: 字体颜色
    :param name: 字体名
    :return:
    '''
    run.font.size = font_size
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = name
    # 设置字体必须要下面2步
    s = run._element
    s.rPr.rFonts.set(qn('w:eastAsia'), name)


def paragraphs_utils(obj):
    for p in obj.paragraphs:
        # 先循环得到单个段落p
        for r in p.runs:
            if word not in r.text:
                # 判断关键字是否存在于段落文本中
                continue
            # print(r.text)
            # print(r.style.name)
            font_size = r.font.size
            bold = r.bold
            color = r.font.color.rgb
            name = u'楷体'
            # 使用关键词切分当前run的文本
            rest = r.text.split(word)
            # 清除当前run的内容
            r.text = ''
            for text in rest[:-1]:
                # 循环切割出来的列表 ['','xxxxxxx']或者['xxxxx','']
                run = p.add_run(text=text)
                set_run(run, font_size, bold, color, name)
                run = p.add_run(word)
                # 重写关键字部分
                set_run(run, font_size, bold, color, name)
                run.font.color.rgb = RGBColor(255, 0, 0)
            run = p.add_run(rest[-1])
            # 在补齐r.text的内容
            set_run(run, font_size, bold, color, name)
    obj.save('标注后的文档.docx')


def table_utils(obj):
    for p in obj.tables:
        # 先循环得到单个表格p
        pd_block = pd.DataFrame(table2list(p))
        # 使用table2list 将table转成列表，然后转成pandas的DateFrame对象
        for rows in range(pd_block.shape[0]):
            # 循环pd_block(DateFrame对象)的行数 -》shape方法得到元祖 为行数和列数
            if rows == 0: continue
            if table_text != pd_block.iloc[rows, 0]: continue
            # 判断关键字是否等于当前表的 rows行0列，否则跳过
            for cols in range(pd_block.shape[1]):
                if cols == 0: continue
                rs = p.cell(rows, cols).paragraphs[0]
                # 此时rows和cols肯定为关键字所在的那行数据，用document对象获取paragraphs取0
                for r in rs.runs:  # paragraphs中有个runs   是个列表
                    font_size = r.font.size
                    bold = r.bold
                    color = r.font.color.rgb
                    name = u'楷体'
                    data = r.text.strip()
                    # 清除当前run的内容
                    r.text = ''
                    run = rs.add_run(data)
                    # 此时要使用paragraphs的add_run方法重写data数据
                    set_run(run, font_size, bold, color, name)
                    run.font.color.rgb = RGBColor(255, 0, 0)
    obj.save('标注后的表格.docx')


for block in iter_block_items(obj):
    if isinstance(block, Paragraph):
        for r in block.runs:
            if word not in r.text:
                continue
            print(r.text)
            print(r.style.name)
            font_size = r.font.size
            bold = r.bold
            color = r.font.color.rgb
            name = u'楷体'
            # 使用关键词切分当前run的文本
            rest = r.text.split(word)
            # 清除当前run的内容
            r.text = ''
            for text in rest[:-1]:
                run = block.add_run(text=text)
                set_run(run, font_size, bold, color, name)
                run = block.add_run(word)
                set_run(run, font_size, bold, color, name)
                run.font.color.rgb = RGBColor(255, 0, 0)
            run = block.add_run(rest[-1])
            set_run(run, font_size, bold, color, name)
    else:
        pd_block = pd.DataFrame(table2list(block))
        # 使用table2list 将table转成列表，然后转成pandas的DateFrame对象
        for rows in range(pd_block.shape[0]):
            # 循环pd_block(DateFrame对象)的行数 -》shape方法得到元祖 为行数和列数
            if rows == 0: continue
            if table_text != pd_block.iloc[rows, 0]: continue
            # 判断关键字是否等于当前表的 rows行0列，否则跳过
            for cols in range(pd_block.shape[1]):
                if cols == 0: continue
                rs = block.cell(rows, cols).paragraphs[0]
                # 此时rows和cols肯定为关键字所在的那行数据，用document对象获取paragraphs取0
                for r in rs.runs:  # paragraphs中有个runs   是个列表
                    font_size = r.font.size
                    bold = r.bold
                    color = r.font.color.rgb
                    name = u'楷体'
                    data = r.text.strip()
                    # 清除当前run的内容
                    r.text = ''
                    run = rs.add_run(data)
                    # 此时要使用paragraphs的add_run方法重写data数据
                    set_run(run, font_size, bold, color, name)
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    
obj.save('段落与表格标注后的文档.docx')